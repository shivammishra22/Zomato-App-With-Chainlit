import os
import re
import pandas as pd
import numpy as np
from docx import Document


class PatientExposureCalculator:
    def __init__(self, ddd_value, country, medicine, place, date):
        self.ddd_value = ddd_value
        self.country = country
        self.medicine = medicine
        self.place = place
        self.date = date
        self.df_combined = pd.DataFrame()
        self.total_exposure_value = 0

    def extract_table_after_text(self, doc, search_text):
        pattern = re.compile(re.escape(search_text), re.IGNORECASE)
        found_index = None

        for i, para in enumerate(doc.paragraphs):
            if pattern.search(para.text):
                found_index = i
                break

        if found_index is None:
            print(f"❌ Text not found: {search_text}")
            return None

        para_counter = 0
        table_counter = 0
        for block in doc.element.body:
            if block.tag.endswith('p'):
                para_counter += 1
            elif block.tag.endswith('tbl'):
                if para_counter > found_index:
                    table = doc.tables[table_counter]
                    break
                table_counter += 1
        else:
            print(f"❌ No table found after: {search_text}")
            return None

        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])

        if len(table_data) > 1 and table_data[0] == table_data[1]:
            table_data.pop(1)

        return table_data

    def save_table_to_excel(self, table_data, excel_path):
        os.makedirs(os.path.dirname(excel_path), exist_ok=True)
        df = pd.DataFrame(table_data[1:], columns=table_data[0])
        df.to_excel(excel_path, index=False)
        print(f"✅ Table saved to: {excel_path}")

    def generate_fallback_doc(self):
        fallback_doc = Document()
        fallback_doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)
        placeholder_text = (
            f"No cumulative and interval patient exposure from marketing experience was available as the MAH "
            f"has not marketed its product {self.medicine} in any country since obtaining initial granting of MA "
            f"till the DLP of this report."
        )
        fallback_doc.add_paragraph(placeholder_text)
        fallback_doc.save(f"{self.medicine}_Exposure.docx")
        print(f"📄 Placeholder Word document saved as '{self.medicine}_Exposure.docx'")

    def calculate_exposure_and_generate_doc(self, excel_path):
        df = pd.read_excel(excel_path, engine='openpyxl')

        df["Strength in mg"] = df["Strength in mg"].astype(str).str.replace("mg", "").str.strip()
        df["Strength in mg"] = pd.to_numeric(df["Strength in mg"], errors='coerce')

        pack_column = next((col for col in ["Pack", "Packs"] if col in df.columns), None)
        if pack_column:
            df[pack_column] = pd.to_numeric(
                df[pack_column].astype(str).str.replace(",", "").str.extract(r'(\d+)')[0],
                errors='coerce'
            ).fillna(0).astype(int)

        if "Pack size" in df.columns:
            pack_size_extracted = df["Pack size"].astype(str).str.extract(r'(\d+)\s*[xX]\s*(\d+)')
            df["Pack size"] = (
                pd.to_numeric(pack_size_extracted[0], errors='coerce').fillna(1).astype(int) *
                pd.to_numeric(pack_size_extracted[1], errors='coerce').fillna(1).astype(int)
            )

        unit_col = "Number of tablets / Capsules/Injections"
        df[unit_col] = df[unit_col].astype(str).str.replace(",", "").str.split(":").str[-1].str.strip()
        df[unit_col] = pd.to_numeric(df[unit_col], errors='coerce')

        df["Delivered quantity (mg)"] = pd.to_numeric(
            df["Delivered quantity (mg)"].astype(str).str.replace(",", ""),
            errors='coerce'
        )

        df.rename(columns={"Product": "Molecule"}, inplace=True)

        df["Sales Figure (mg) or period/Volume of sales (in mg)"] = df[unit_col] * df["Strength in mg"]
        df["Patients Exposure (PTY) for period"] = (
            df["Sales Figure (mg) or period/Volume of sales (in mg)"] / (self.ddd_value * 365)
        ).round(0)

        df_country = df[df["Country"] == self.country].copy()
        df_non_country = df[df["Country"] != self.country].copy()

        def create_clean_total_row(dataframe):
            total = dataframe["Patients Exposure (PTY) for period"].sum()
            total_row = {col: "" for col in dataframe.columns}
            total_row["Country"] = "Total"
            total_row["Patients Exposure (PTY) for period"] = int(total)
            return pd.DataFrame([total_row])

        df_country = pd.concat([df_country, create_clean_total_row(df_country)], ignore_index=True)
        df_non_country = pd.concat([df_non_country, create_clean_total_row(df_non_country)], ignore_index=True)

        df_country.fillna("", inplace=True)
        df_non_country.fillna("", inplace=True)

        sa_total = df_country[df_country["Country"] == "Total"]["Patients Exposure (PTY) for period"].values[0]
        if sa_total == 0:
            self.generate_fallback_doc()
            return

        doc = Document()
        doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)

        non_country_total = df_non_country[df_non_country["Country"] == "Total"]["Patients Exposure (PTY) for period"].values[0]
        total_exposure = sa_total + non_country_total
        self.total_exposure_value = int(total_exposure)

        summary_text = (
            f"The MAH obtained initial MA for their generic formulation of {self.medicine} in {self.place} on {self.date}.\n"
            f"The post-authorization exposure to {self.medicine} during the cumulative period was {int(total_exposure)} patients "
            f"({self.place}: {int(sa_total)} and Non {self.place}: {int(non_country_total)}) treatment days approximately and presented in Table 3."
        )
        doc.add_paragraph(summary_text)

        def add_table_with_data(doc, dataframe, title):
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=1, cols=len(dataframe.columns))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(dataframe.columns):
                hdr_cells[i].text = str(col_name)
            for _, row in dataframe.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)

        add_table_with_data(doc, df_country, f"{self.country} ")
        add_table_with_data(doc, df_non_country, f"Non-{self.country} ")

        doc.save(f"{self.medicine}_Exposure.docx")
        print(f"✅ Word document saved as '{self.medicine}_Exposure.docx'")

        self.df_combined = pd.concat([df_country, df_non_country], ignore_index=True)

    def get_total_patient_exposure(self):
        return self.total_exposure_value

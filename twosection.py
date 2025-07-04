import os
import re
import pandas as pd
import numpy as np
from docx import Document

def extract_table_after_text(doc, search_text):
    pattern = re.compile(re.escape(search_text), re.IGNORECASE)
    found_index = None

    for i, para in enumerate(doc.paragraphs):
        if pattern.search(para.text):
            found_index = i
            break

    if found_index is None:
        print(f"❌ Text not found: {search_text}")
        return None  # Changed from [] to None for clarity

    para_counter = 0
    table_counter = 0
    for block in doc.element.body:
        if block.tag.endswith('p'):
            para_counter += 1
        elif block.tag.endswith('tbl'):
            if para_counter > found_index:
                table = doc.tables[table_counter]
                break
            table_counter += 1
    else:
        print(f"❌ No table found after: {search_text}")
        return None

    table_data = []
    for row in table.rows:
        table_data.append([cell.text.strip() for cell in row.cells])

    if len(table_data) > 1 and table_data[0] == table_data[1]:
        table_data.pop(1)

    return table_data

def save_table_to_excel(table_data, excel_path):
    os.makedirs(os.path.dirname(excel_path), exist_ok=True)
    df = pd.DataFrame(table_data[1:], columns=table_data[0])
    df.to_excel(excel_path, index=False)
    print(f"✅ Table saved to: {excel_path}")

def calculate_exposure_and_generate_doc(excel_path, ddd_value, country_name, medicine, place, date):
    df = pd.read_excel(excel_path, engine='openpyxl')

    df["Strength in mg"] = df["Strength in mg"].astype(str).str.replace("mg", "").str.strip()
    df["Strength in mg"] = pd.to_numeric(df["Strength in mg"], errors='coerce')

    pack_column = next((col for col in ["Pack", "Packs"] if col in df.columns), None)
    if pack_column:
        df[pack_column] = pd.to_numeric(df[pack_column].astype(str).str.replace(",", "").str.extract(r'(\d+)')[0], errors='coerce').fillna(0).astype(int)

    if "Pack size" in df.columns:
        pack_size_extracted = df["Pack size"].astype(str).str.extract(r'(\d+)\s*[xX]\s*(\d+)')
        df["Pack size"] = pd.to_numeric(pack_size_extracted[0], errors='coerce').fillna(1).astype(int) * pd.to_numeric(pack_size_extracted[1], errors='coerce').fillna(1).astype(int)

    unit_col = "Number of tablets / Capsules/Injections"
    df[unit_col] = df[unit_col].astype(str).str.replace(",", "").str.split(":").str[-1].str.strip()
    df[unit_col] = pd.to_numeric(df[unit_col], errors='coerce')

    df["Delivered quantity (mg)"] = pd.to_numeric(df["Delivered quantity (mg)"].astype(str).str.replace(",", ""), errors='coerce')
    df.rename(columns={"Product": "Molecule"}, inplace=True)

    df["Sales Figure (mg) or period/Volume of sales (in mg)"] = df[unit_col] * df["Strength in mg"]
    df["Patients Exposure (PTY) for period"] = df["Sales Figure (mg) or period/Volume of sales (in mg)"] / (ddd_value * 365)
    df["Patients Exposure (PTY) for period"] = df["Patients Exposure (PTY) for period"].round(0)

    df_country = df[df["Country"] == country_name].copy()
    df_non_country = df[df["Country"] != country_name].copy()

    def create_clean_total_row(dataframe):
        total = dataframe["Patients Exposure (PTY) for period"].sum()
        total_row = {col: "" for col in dataframe.columns}
        total_row["Country"] = "Total"
        total_row["Patients Exposure (PTY) for period"] = int(total)
        return pd.DataFrame([total_row])

    df_country = pd.concat([df_country, create_clean_total_row(df_country)], ignore_index=True)
    df_non_country = pd.concat([df_non_country, create_clean_total_row(df_non_country)], ignore_index=True)

    df_country.fillna("", inplace=True)
    df_non_country.fillna("", inplace=True)

    doc = Document()
    doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)

    country_total_exposure = df_country[df_country["Country"] == "Total"]["Patients Exposure (PTY) for period"].values[0]
    non_country_total_exposure = df_non_country[df_non_country["Country"] == "Total"]["Patients Exposure (PTY) for period"].values[0]
    total_exposure = country_total_exposure + non_country_total_exposure

    summary_text = (
        f"The MAH obtained initial MA for their generic formulation of {medicine} in {place} on {date}.\n"
        f"The post-authorization exposure to {medicine} during the cumulative period was {int(total_exposure)} patients "
        f"({place}: {int(country_total_exposure)} and Non {place}: {int(non_country_total_exposure)}) treatment days approximately and presented in Table 3."
    )
    doc.add_paragraph(summary_text)

    def add_table_with_data(doc, dataframe, title):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(dataframe.columns):
            hdr_cells[i].text = str(col_name)
        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    add_table_with_data(doc, df_country, f"                                                      {country_name} ")
    add_table_with_data(doc, df_non_country, f"                                                      Non-{country_name} ")

    doc.save("Esomeprazole_Exposure.docx")
    print("✅ Word document saved as 'Esomeprazole_Exposure.docx'")

# === MAIN EXECUTION ===

if __name__ == "__main__":
    docx_path = r"C:\Users\shivam.mishra2\Downloads\Data request form.docx"
    search_text = "Cumulative sales data sale required"
    excel_output_path = r"C:\Users\shivam.mishra2\Downloads\New_Psur_File\marketing_exposure_tables.xlsx"

    ddd_value = 10
    country = "South Africa"
    medicine = "Esomeprazole"
    place = "South Africa"
    date = "2020-01-01"

    try:
        doc = Document(docx_path)
        table_data = extract_table_after_text(doc, search_text)
        if table_data:
            save_table_to_excel(table_data, excel_output_path)
            calculate_exposure_and_generate_doc(excel_output_path, ddd_value, country, medicine, place, date)
        else:
            raise ValueError("Search text or table not found.")
    except Exception as e:
        print(f"⚠️ Error: {e}")
        fallback_doc = Document()
        fallback_doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)
        placeholder_text = (
            f"No cumulative and interval patient exposure from marketing experience was available as the MAH "
            f"has not marketed its product {medicine} in any country since obtaining initial granting of MA "
            f"till the DLP of this report."
        )
        fallback_doc.add_paragraph(placeholder_text)
        fallback_doc.save("Esomeprazole_Exposure.docx")
        print("📄 Placeholder Word document saved as 'Esomeprazole_Exposure.docx'")

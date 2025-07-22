#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import pandas as pd
import numpy as np
from docx import Document
import requests
from bs4 import BeautifulSoup
import urllib3

# =========================================================
# === 0. CONFIG / CONSTANTS ===============================
# =========================================================

DOCX_PATH = r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Olanzapine PSUR_South Africa_29-Sep-17 to 31-Mar-25\Draft\DRA\Data request form_olanzapine.docx"
SEARCH_TEXT = "Cumulative sales data sale required"
EXCEL_OUTPUT_PATH = r"C:\Users\shivam.mishra2\Downloads\New_Psur_File\marketing_exposure_tables.xlsx"
DDD_EXCEL_PATH = r"drug_code_map_with_ddd.xlsx"

COUNTRY = "South Africa"
MEDICINE = "Olanzapine"
PLACE = "South Africa"
DATE = "2020-01-01"

# === NEW / UPDATED ===
# Map of product names to dosage forms
PRODUCT_DOSAGE_MAP = {
    "Esomeprazole": "Gastro-resistant",
    "Zipola 5": "Film coated Tablet",
    "Zipola 10": "Film coated Tablet",
    "Jubilonz OD10": "Oro dispersible tablet",
    "Jubilonz OD5": "Oro dispersible tablet",
    "SCHIZOLANZ": "Oro dispersible tablet",
    "Olanzapine film coated tablets": "Film coated Tablet",
    "Olanzapine": "Film coated Tablet"  # add generic too, just in case
}

# =========================================================
# === 1. UTILITIES ========================================
# =========================================================

def extract_table_after_text(doc, search_text):
    pattern = re.compile(re.escape(search_text), re.IGNORECASE)
    found_index = None

    for i, para in enumerate(doc.paragraphs):
        if pattern.search(para.text):
            found_index = i
            break

    if found_index is None:
        print(f"❌ Text not found: {search_text}")
        return None

    para_counter = 0
    table_counter = 0
    for block in doc.element.body:
        if block.tag.endswith('p'):
            para_counter += 1
        elif block.tag.endswith('tbl'):
            if para_counter > found_index:
                table = doc.tables[table_counter]
                break
            table_counter += 1
    else:
        print(f"❌ No table found after: {search_text}")
        return None

    table_data = []
    for row in table.rows:
        table_data.append([cell.text.strip() for cell in row.cells])

    # Remove duplicate header row if Word duplicated it
    if len(table_data) > 1 and table_data[0] == table_data[1]:
        table_data.pop(1)

    return table_data


def save_table_to_excel(table_data, excel_path):
    os.makedirs(os.path.dirname(excel_path), exist_ok=True)
    df = pd.DataFrame(table_data[1:], columns=table_data[0])
    df.to_excel(excel_path, index=False)
    print(f"✅ Table saved to: {excel_path}")


def generate_fallback_doc(medicine):
    fallback_doc = Document()
    fallback_doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)
    placeholder_text = (
        f"No cumulative and interval patient exposure from marketing experience was available as the MAH "
        f"has not marketed its product {medicine} in any country since obtaining initial granting of MA "
        f"till the DLP of this report."
    )
    fallback_doc.add_paragraph(placeholder_text)
    fallback_doc.save(f"{medicine}_Exposure.docx")
    print(f"📄 Placeholder Word document saved as '{medicine}_Exposure.docx'")


def add_table_with_data(doc, dataframe, title):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(dataframe.columns):
        hdr_cells[i].text = str(col_name)
    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = "" if pd.isna(item) else str(item)


def create_clean_total_row(dataframe, total_col="Patients Exposure (PTY) for period"):
    total = dataframe[total_col].sum(numeric_only=True)
    total_row = {col: "" for col in dataframe.columns}
    total_row["Country"] = "Total"
    total_row[total_col] = int(total)
    return pd.DataFrame([total_row])


def map_dosage(product_name):
    """Return dosage form string if any key in PRODUCT_DOSAGE_MAP is substring of product_name (case-insensitive)."""
    if pd.isna(product_name):
        return ""
    name = str(product_name).lower()
    for key, val in PRODUCT_DOSAGE_MAP.items():
        if key.lower() in name:
            return val
    return ""


def add_dosage_column(df):
    """
    Safely add Dosage Form column. Tries 'Product' first, then 'Molecule'.
    """
    col_to_use = None
    for c in ["Product", "Molecule"]:
        if c in df.columns:
            col_to_use = c
            break

    if col_to_use is None:
        print("⚠️ Neither 'Product' nor 'Molecule' column found. Skipping dosage mapping.")
        return df

    df["Dosage Form (Units)"] = df[col_to_use].apply(map_dosage)
    return df


# =========================================================
# === 2. CORE PROCESS =====================================
# =========================================================

def calculate_exposure_and_generate_doc(excel_path, ddd_value, country_name, medicine, place, date):
    df = pd.read_excel(excel_path, engine='openpyxl')

    # === NEW / UPDATED ===
    # Ensure dosage form column exists
    df = add_dosage_column(df)

    # Clean / convert columns
    if "Strength in mg" in df.columns:
        df["Strength in mg"] = df["Strength in mg"].astype(str).str.replace("mg", "", regex=False).str.strip()
        df["Strength in mg"] = pd.to_numeric(df["Strength in mg"], errors='coerce')

    pack_column = next((col for col in ["Pack", "Packs"] if col in df.columns), None)
    if pack_column:
        df[pack_column] = (
            df[pack_column]
            .astype(str)
            .str.replace(",", "", regex=False)
            .str.extract(r'(\d+)')[0]
        )
        df[pack_column] = pd.to_numeric(df[pack_column], errors='coerce').fillna(0).astype(int)

    if "Pack size" in df.columns:
        pack_size_extracted = df["Pack size"].astype(str).str.extract(r'(\d+)\s*[xX]\s*(\d+)')
        df["Pack size"] = (
            pd.to_numeric(pack_size_extracted[0], errors='coerce').fillna(1).astype(int)
            * pd.to_numeric(pack_size_extracted[1], errors='coerce').fillna(1).astype(int)
        )

    unit_col = "Number of tablets / Capsules/Injections"
    if unit_col in df.columns:
        df[unit_col] = (
            df[unit_col].astype(str)
            .str.replace(",", "", regex=False)
            .str.split(":").str[-1]
            .str.strip()
        )
        df[unit_col] = pd.to_numeric(df[unit_col], errors='coerce')

    if "Delivered quantity (mg)" in df.columns:
        df["Delivered quantity (mg)"] = pd.to_numeric(
            df["Delivered quantity (mg)"].astype(str).str.replace(",", "", regex=False),
            errors='coerce'
        )

    # Rename Product -> Molecule once (if Product exists)
    if "Product" in df.columns and "Molecule" not in df.columns:
        df.rename(columns={"Product": "Molecule"}, inplace=True)

    # Add DDD column
    df['DDD*'] = f'{int(ddd_value)} mg'

    # Calculate Sales & Exposure
    if unit_col in df.columns and "Strength in mg" in df.columns:
        df["Sales Figure (mg) or period/Volume of sales (in mg)"] = df[unit_col] * df["Strength in mg"]
    else:
        df["Sales Figure (mg) or period/Volume of sales (in mg)"] = np.nan

    df["Patients Exposure (PTY) for period"] = (
        df["Sales Figure (mg) or period/Volume of sales (in mg)"] / (ddd_value * 365)
    )
    df["Patients Exposure (PTY) for period"] = df["Patients Exposure (PTY) for period"].round(0)

    # Split by country
    if "Country" not in df.columns:
        print("⚠️ 'Country' column missing in Excel. Will treat all as Non-country.")
        df["Country"] = "Unknown"

    df_country = df[df["Country"] == country_name].copy()
    df_non_country = df[df["Country"] != country_name].copy()

    # Totals
    df_country_total_row = create_clean_total_row(df_country)
    df_non_country_total_row = create_clean_total_row(df_non_country)

    df_country = pd.concat([df_country, df_country_total_row], ignore_index=True)
    df_non_country = pd.concat([df_non_country, df_non_country_total_row], ignore_index=True)

    df_country.fillna("", inplace=True)
    df_non_country.fillna("", inplace=True)

    sa_total = df_country_total_row["Patients Exposure (PTY) for period"].values[0]
    non_sa_total = df_non_country_total_row["Patients Exposure (PTY) for period"].values[0]
    combined_total = int(sa_total) + int(non_sa_total)

    print(f"📊 Combined Total Exposure: {combined_total}")

    if sa_total == 0:
        generate_fallback_doc(medicine)
        return

    # Generate Word doc
    doc = Document()
    doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)

    summary_text = (
        f"The MAH obtained initial MA for their generic formulation of {medicine} in {place} on {date}.\n"
        f"The post-authorization exposure to {medicine} during the cumulative period was {combined_total} patients "
        f"({place}: {int(sa_total)} and Non {place}: {int(non_sa_total)}) treatment days approximately and presented in Table 3."
    )
    doc.add_paragraph(summary_text)

    # === NEW / UPDATED === Ensure Dosage Form col appears in Word too
    add_table_with_data(doc, df_country, f"                                                      {country_name} ")
    add_table_with_data(doc, df_non_country, f"                                                      Non-{country_name} ")

    out_doc_name = f"{medicine}_Exposure.docx"
    doc.save(out_doc_name)
    print(f"✅ Word document saved as '{out_doc_name}'")


# =========================================================
# === 3. DDD FALLBACK FETCH ===============================
# =========================================================

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

def fetch_ddd_fallback(medicine, code):
    if pd.isna(code):
        return np.nan
    url = f"https://atcddd.fhi.no/atc_ddd_index/?code={code}"
    try:
        response = requests.get(url, verify=False, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        ddd_values = soup.find_all("td", align="right")
        for td in ddd_values:
            value = td.get_text(strip=True)
            if value.replace('.', '', 1).isdigit():
                print(f"Fetched fallback DDD for {medicine} ({code}): {value}")
                return float(value)
        print(f"No DDD value found on fallback site for {medicine} ({code})")
        return np.nan
    except Exception as e:
        print(f"Error during fallback DDD fetch for {medicine} ({code}): {e}")
        return np.nan


# =========================================================
# === 4. MAIN =============================================
# =========================================================

if __name__ == "__main__":
    # Read the DDD Excel (if available)
    try:
        ddd_df = pd.read_excel(DDD_EXCEL_PATH)
    except Exception as e:
        print(f"⚠️ Could not read DDD Excel: {e}")
        ddd_df = pd.DataFrame()

    try:
        if not os.path.exists(DOCX_PATH):
            raise FileNotFoundError(f"File not found: {DOCX_PATH}")

        doc = Document(DOCX_PATH)
        table_data = extract_table_after_text(doc, SEARCH_TEXT)

        # --- DDD VALUE LOGIC START ---
        ddd_value = np.nan
        if not ddd_df.empty and "Drug Name" in ddd_df.columns:
            ddd_row = ddd_df[ddd_df["Drug Name"].str.lower() == MEDICINE.lower()]
        else:
            ddd_row = pd.DataFrame()

        if not ddd_row.empty and "DDD Value" in ddd_row.columns and not pd.isna(ddd_row.iloc[0]["DDD Value"]):
            ddd_value = ddd_row.iloc[0]["DDD Value"]
            print(f"✅ DDD value found in Excel: {ddd_value}")
        else:
            code = np.nan
            if not ddd_row.empty and "Drug Code" in ddd_row.columns:
                code = ddd_row.iloc[0]["Drug Code"]
            elif "Drug Code" in getattr(ddd_df, "columns", []):
                possible = ddd_df[ddd_df["Drug Name"].str.lower().str.contains(MEDICINE.lower(), na=False)]
                if not possible.empty:
                    code = possible.iloc[0]["Drug Code"]

            ddd_value = fetch_ddd_fallback(MEDICINE, code)
            if pd.isna(ddd_value):
                print("❌ DDD value not found anywhere. Will use fallback document.")
        # --- DDD VALUE LOGIC END ---

        if table_data and not pd.isna(ddd_value):
            save_table_to_excel(table_data, EXCEL_OUTPUT_PATH)

            # === NEW / UPDATED ===
            # After saving initial excel, reopen & add dosage column (if not already inside calc fn)
            calculate_exposure_and_generate_doc(EXCEL_OUTPUT_PATH, ddd_value, COUNTRY, MEDICINE, PLACE, DATE)

            # Overwrite the Excel with dosage column if needed (calc function already added it)
            try:
                final_df = pd.read_excel(EXCEL_OUTPUT_PATH, engine='openpyxl')
                final_df = add_dosage_column(final_df)  # idempotent
                final_df.to_excel(EXCEL_OUTPUT_PATH, index=False)
                print(f"✅ Updated Excel saved with 'Dosage Form (Units)' column at: {EXCEL_OUTPUT_PATH}")
            except Exception as e:
                print(f"❌ Error saving updated Excel: {e}")

        else:
            print("⚠️ Table not found after search text or DDD is missing. Generating fallback document.")
            generate_fallback_doc(MEDICINE)

    except Exception as e:
        print(f"⚠️ Error: {e}")
        generate_fallback_doc(MEDICINE)
                

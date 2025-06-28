from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentStyling:
    """Centralized styling configuration"""
    
    # Font settings
    FONT_NAME = 'Times New Roman'
    
    # Font sizes
    HEADING_FONT_SIZE = Pt(14)
    SUBHEADING_FONT_SIZE = Pt(12)
    CONTENT_FONT_SIZE = Pt(11)
    
    # Colors
    BLACK_COLOR = RGBColor(0, 0, 0)
    
    # Alignment
    JUSTIFY_ALIGNMENT = WD_ALIGN_PARAGRAPH.JUSTIFY
    LEFT_ALIGNMENT = WD_ALIGN_PARAGRAPH.LEFT
    CENTER_ALIGNMENT = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT_ALIGNMENT = WD_ALIGN_PARAGRAPH.RIGHT
    
    @staticmethod
    def apply_heading_style(run):
        """Apply heading style (size 14, bold, underlined, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.HEADING_FONT_SIZE
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_heading_number_style(run):
        """Apply heading number style (size 14, bold, NO underline, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.HEADING_FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_heading_text_style(run):
        """Apply heading text style (size 14, bold, underlined, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.HEADING_FONT_SIZE
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_subheading_style(run):
        """Apply subheading style (size 12, bold, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.SUBHEADING_FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_subheading_number_style(run):
        """Apply subheading number style (size 12, bold, NO underline, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.SUBHEADING_FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_subheading_text_style(run):
        """Apply subheading text style (size 12, bold, underlined, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.SUBHEADING_FONT_SIZE
        run.font.bold = True
        run.font.underline = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_content_style(run):
        """Apply content style (size 11, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.CONTENT_FONT_SIZE
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def apply_bold_content_style(run):
        """Apply bold content style (size 11, bold, Times New Roman)"""
        run.font.name = DocumentStyling.FONT_NAME
        run.font.size = DocumentStyling.CONTENT_FONT_SIZE
        run.font.bold = True
        run.font.color.rgb = DocumentStyling.BLACK_COLOR
    
    @staticmethod
    def set_paragraph_alignment(paragraph, alignment):
        """Set paragraph alignment"""
        paragraph.alignment = alignment
    
    @staticmethod
    def apply_table_header_style(cell):
        """Apply table header styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                DocumentStyling.apply_bold_content_style(run)
    
    @staticmethod
    def create_split_heading(paragraph, heading_text):
        """Create a heading with number not underlined and text underlined"""
        import re
        
        # Split heading into number and text parts
        match = re.match(r'^(\d+\.?\d*\.?\s*)(.*)', heading_text)
        if match:
            number_part = match.group(1)
            text_part = match.group(2)
            
            # Add number part (bold, no underline)
            number_run = paragraph.add_run(number_part)
            DocumentStyling.apply_heading_number_style(number_run)
            
            # Add text part (bold, underlined)
            text_run = paragraph.add_run(text_part)
            DocumentStyling.apply_heading_text_style(text_run)
        else:
            # If no number found, just underline the whole text
            run = paragraph.add_run(heading_text)
            DocumentStyling.apply_heading_text_style(run)
    
    @staticmethod
    def create_split_subheading(paragraph, subheading_text):
        """Create a subheading with number not underlined and text underlined"""
        import re
        
        # Split subheading into number and text parts
        match = re.match(r'^(\d+\.?\d*\.?\s*)(.*)', subheading_text)
        if match:
            number_part = match.group(1)
            text_part = match.group(2)
            
            # Add number part (bold, no underline)
            number_run = paragraph.add_run(number_part)
            DocumentStyling.apply_subheading_number_style(number_run)
            
            # Add text part (bold, underlined)
            text_run = paragraph.add_run(text_part)
            DocumentStyling.apply_subheading_text_style(text_run)
        else:
            # If no number found, just underline the whole text
            run = paragraph.add_run(subheading_text)
            DocumentStyling.apply_subheading_text_style(run)

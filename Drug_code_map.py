#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os, re, urllib3, requests
import pandas as pd
import numpy as np
from docx import Document
from bs4 import BeautifulSoup

# =========================================================
# 0. CONFIG
# =========================================================
DOCX_PATH = r"C:\Users\shivam.mishra2\Downloads\ALL_PSUR_File\PSUR_all _Data\Olanzapine PSUR_South Africa_29-Sep-17 to 31-Mar-25\Draft\DRA\Data request form_olanzapine.docx"
SEARCH_TEXT = "Cumulative sales data sale required"
EXCEL_OUTPUT_PATH = r"C:\Users\shivam.mishra2\Downloads\New_Psur_File\marketing_exposure_tables.xlsx"
DDD_EXCEL_PATH = r"drug_code_map_with_ddd.xlsx"

COUNTRY  = "South Africa"
MEDICINE = "Olanzapine"
PLACE    = "South Africa"
DATE     = "2020-01-01"

PRODUCT_DOSAGE_MAP = {
    "Esomeprazole": "Gastro-resistant",
    "Zipola 5": "Film coated Tablet",
    "Zipola 10": "Film coated Tablet",
    "Jubilonz OD10": "Oro dispersible tablet",
    "Jubilonz OD5": "Oro dispersible tablet",
    "SCHIZOLANZ": "Oro dispersible tablet",
    "Olanzapine film coated tablets": "Film coated Tablet",
    "Olanzapine": "Film coated Tablet"
}

# =========================================================
# 1. UTILITIES
# =========================================================
def extract_table_after_text(doc, search_text):
    pattern = re.compile(re.escape(search_text), re.IGNORECASE)
    found_index = None
    for i, para in enumerate(doc.paragraphs):
        if pattern.search(para.text):
            found_index = i
            break
    if found_index is None:
        print(f"❌ Text not found: {search_text}")
        return None

    para_counter = 0
    table_counter = 0
    for block in doc.element.body:
        if block.tag.endswith('p'):
            para_counter += 1
        elif block.tag.endswith('tbl'):
            if para_counter > found_index:
                table = doc.tables[table_counter]
                break
            table_counter += 1
    else:
        print(f"❌ No table found after: {search_text}")
        return None

    table_data = []
    for row in table.rows:
        table_data.append([cell.text.strip() for cell in row.cells])

    if len(table_data) > 1 and table_data[0] == table_data[1]:
        table_data.pop(1)
    return table_data


def save_table_to_excel(table_data, excel_path):
    os.makedirs(os.path.dirname(excel_path), exist_ok=True)
    df = pd.DataFrame(table_data[1:], columns=table_data[0])
    df.to_excel(excel_path, index=False)
    print(f"✅ Table saved to: {excel_path}")


def generate_fallback_doc(medicine):
    doc = Document()
    doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)
    doc.add_paragraph(
        f"No cumulative and interval patient exposure from marketing experience was available as the MAH "
        f"has not marketed its product {medicine} in any country since obtaining initial granting of MA "
        f"till the DLP of this report."
    )
    out = f"{medicine}_Exposure.docx"
    doc.save(out)
    print(f"📄 Placeholder Word document saved as '{out}'")


def map_dosage(product_name):
    if pd.isna(product_name):
        return ""
    name = str(product_name).lower()
    for k, v in PRODUCT_DOSAGE_MAP.items():
        if k.lower() in name:
            return v
    return ""


def add_dosage_column(df):
    src_col = None
    for c in ["Product", "Molecule"]:
        if c in df.columns:
            src_col = c
            break
    if src_col is None:
        print("⚠️ Neither 'Product' nor 'Molecule' found. Skipping dosage mapping.")
        return df
    df["Dosage Form (Units)"] = df[src_col].apply(map_dosage)
    return df


def nice_int(x):
    try:
        return f"{int(float(str(x).replace(',', ''))):,}"
    except Exception:
        return x


def format_df_for_report(df: pd.DataFrame, ddd_value: float) -> pd.DataFrame:
    """Rename + reorder columns to: Country | Molecule | Dosage Form (Units) | Formulation Strength | DDD* |
       Pack size | Sales figure (units)/Quantity Sold | Sales Figure (mg) or period/Volume of sales (in mg) |
       Patients Exposure (PTY) for period
    """
    df = add_dosage_column(df)

    # Formulation Strength from 'Strength in mg'
    if "Strength in mg" in df.columns:
        df["Formulation Strength"] = df["Strength in mg"].apply(
            lambda v: f"{int(v)} mg" if pd.notna(v) else "")

    df["DDD*"] = f"{int(ddd_value)} mg"

    # Sales figure units
    unit_col = "Number of tablets / Capsules/Injections"
    if unit_col in df.columns:
        df["Sales figure (units)/Quantity Sold"] = df[unit_col]

    # final rename not necessary if already same
    desired = [
        "Country",
        "Molecule",
        "Dosage Form (Units)",
        "Formulation Strength",
        "DDD*",
        "Pack size",
        "Sales figure (units)/Quantity Sold",
        "Sales Figure (mg) or period/Volume of sales (in mg)",
        "Patients Exposure (PTY) for period",
    ]

    present = [c for c in desired if c in df.columns]
    others  = [c for c in df.columns if c not in present]
    df = df[present + others]

    # number formatting
    for col in ["Sales figure (units)/Quantity Sold",
                "Sales Figure (mg) or period/Volume of sales (in mg)",
                "Patients Exposure (PTY) for period"]:
        if col in df.columns:
            df[col] = df[col].apply(nice_int)

    return df


def add_table_with_data(doc: Document, dataframe: pd.DataFrame, title: str):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(dataframe.columns):
        hdr_cells[i].text = str(col_name)

    for _, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = "" if pd.isna(item) else str(item)


def create_clean_total_row(df_, total_col="Patients Exposure (PTY) for period"):
    numeric = pd.to_numeric(df_[total_col].astype(str).str.replace(',', ''), errors='coerce')
    total = int(numeric.sum()) if pd.notna(numeric.sum()) else 0
    blank = {c: "" for c in df_.columns}
    blank["Country"] = "Total"
    blank[total_col] = total
    return pd.DataFrame([blank])

# =========================================================
# 2. CORE
# =========================================================
def calculate_exposure_and_generate_doc(excel_path, ddd_value, country_name, medicine, place, date):
    df = pd.read_excel(excel_path, engine='openpyxl')

    # Clean & numeric prep
    if "Strength in mg" in df.columns:
        df["Strength in mg"] = df["Strength in mg"].astype(str).str.replace("mg", "", regex=False).str.strip()
        df["Strength in mg"] = pd.to_numeric(df["Strength in mg"], errors='coerce')

    unit_col = "Number of tablets / Capsules/Injections"
    if unit_col in df.columns:
        df[unit_col] = (df[unit_col].astype(str)
                        .str.replace(",", "", regex=False)
                        .str.split(":").str[-1]
                        .str.strip())
        df[unit_col] = pd.to_numeric(df[unit_col], errors='coerce')

    if "Sales Figure (mg) or period/Volume of sales (in mg)" not in df.columns:
        if unit_col in df.columns and "Strength in mg" in df.columns:
            df["Sales Figure (mg) or period/Volume of sales (in mg)"] = df[unit_col] * df["Strength in mg"]

    df["Patients Exposure (PTY) for period"] = (
        df["Sales Figure (mg) or period/Volume of sales (in mg)"] / (ddd_value * 365)
    ).round(0)

    if "Product" in df.columns and "Molecule" not in df.columns:
        df.rename(columns={"Product": "Molecule"}, inplace=True)

    if "Country" not in df.columns:
        print("⚠️ 'Country' column missing. Setting all to 'Unknown'")
        df["Country"] = "Unknown"

    df_sa  = df[df["Country"] == country_name].copy()
    df_non = df[df["Country"] != country_name].copy()

    # Format & totals
    df_sa_fmt  = format_df_for_report(df_sa,  ddd_value)
    df_non_fmt = format_df_for_report(df_non, ddd_value)

    df_sa_tot  = create_clean_total_row(df_sa_fmt)
    df_non_tot = create_clean_total_row(df_non_fmt)

    df_sa_fmt  = pd.concat([df_sa_fmt,  df_sa_tot],  ignore_index=True)
    df_non_fmt = pd.concat([df_non_fmt, df_non_tot], ignore_index=True)

    sa_total     = int(df_sa_tot["Patients Exposure (PTY) for period"].iloc[0])
    non_sa_total = int(df_non_tot["Patients Exposure (PTY) for period"].iloc[0])
    combined     = sa_total + non_sa_total
    print(f"📊 Combined Total Exposure: {combined}")

    if sa_total == 0:
        generate_fallback_doc(medicine)
        return

    # Word
    doc = Document()
    doc.add_heading("5.3 Cumulative and Interval Patient Exposure from Marketing Experience", level=1)
    doc.add_paragraph(
        f"The MAH obtained initial MA for their generic formulation of {medicine} in {place} on {date}.\n"
        f"The post-authorization exposure to {medicine} during the cumulative period was {combined} patients "
        f"({place}: {sa_total} and Non {place}: {non_sa_total}) treatment days approximately and presented in Table 3."
    )

    add_table_with_data(doc, df_sa_fmt,  f"                                                      {country_name} ")
    add_table_with_data(doc, df_non_fmt, f"                                                      Non-{country_name} ")

    out_doc = f"{medicine}_Exposure.docx"
    doc.save(out_doc)
    print(f"✅ Word document saved as '{out_doc}'")

    # Overwrite Excel with formatted combined data
    final_df = pd.concat([
        df_sa_fmt.assign(_Section="SA"),
        df_non_fmt.assign(_Section="Non-SA")
    ], ignore_index=True)
    final_df.to_excel(EXCEL_OUTPUT_PATH, index=False)
    print(f"✅ Final Excel overwritten with formatted columns: {EXCEL_OUTPUT_PATH}")

# =========================================================
# 3. DDD FALLBACK
# =========================================================
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

def fetch_ddd_fallback(medicine, code):
    if pd.isna(code):
        return np.nan
    url = f"https://atcddd.fhi.no/atc_ddd_index/?code={code}"
    try:
        r = requests.get(url, verify=False, timeout=10)
        r.raise_for_status()
        soup = BeautifulSoup(r.content, "html.parser")
        for td in soup.find_all("td", align="right"):
            val = td.get_text(strip=True)
            if val.replace('.', '', 1).isdigit():
                print(f"Fetched fallback DDD for {medicine} ({code}): {val}")
                return float(val)
        print(f"No DDD value found for {medicine} ({code})")
        return np.nan
    except Exception as e:
        print(f"Error during fallback DDD fetch for {medicine} ({code}): {e}")
        return np.nan

# =========================================================
# 4. MAIN
# =========================================================
if __name__ == "__main__":
    try:
        ddd_df = pd.read_excel(DDD_EXCEL_PATH)
    except Exception as e:
        print(f"⚠️ Could not read DDD Excel: {e}")
        ddd_df = pd.DataFrame()

    try:
        if not os.path.exists(DOCX_PATH):
            raise FileNotFoundError(f"File not found: {DOCX_PATH}")

        doc = Document(DOCX_PATH)
        table_data = extract_table_after_text(doc, SEARCH_TEXT)

        # --- DDD logic ---
        ddd_value = np.nan
        if not ddd_df.empty and "Drug Name" in ddd_df.columns:
            ddd_row = ddd_df[ddd_df["Drug Name"].str.lower() == MEDICINE.lower()]
        else:
            ddd_row = pd.DataFrame()

        if (not ddd_row.empty and
            "DDD Value" in ddd_row.columns and
            not pd.isna(ddd_row.iloc[0]["DDD Value"])):
            ddd_value = ddd_row.iloc[0]["DDD Value"]
            print(f"✅ DDD value found in Excel: {ddd_value}")
        else:
            code = np.nan
            if not ddd_row.empty and "Drug Code" in ddd_row.columns:
                code = ddd_row.iloc[0]["Drug Code"]
            elif "Drug Code" in getattr(ddd_df, "columns", []):
                poss = ddd_df[ddd_df["Drug Name"].str.lower().str.contains(MEDICINE.lower(), na=False)]
                if not poss.empty:
                    code = poss.iloc[0]["Drug Code"]
            ddd_value = fetch_ddd_fallback(MEDICINE, code)
            if pd.isna(ddd_value):
                print("❌ DDD value not found anywhere. Will use fallback document.")

        if table_data and not pd.isna(ddd_value):
            save_table_to_excel(table_data, EXCEL_OUTPUT_PATH)
            calculate_exposure_and_generate_doc(EXCEL_OUTPUT_PATH, ddd_value, COUNTRY, MEDICINE, PLACE, DATE)
        else:
            print("⚠️ Table not found or DDD missing. Generating fallback document.")
            generate_fallback_doc(MEDICINE)

    except Exception as e:
        print(f"⚠️ Error: {e}")
        generate_fallback_doc(MEDICINE)
    
